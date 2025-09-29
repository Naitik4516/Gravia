import os
import sys
from pathlib import Path
from typing import List, Optional
import logging
import magic
import asyncio

# Third-party imports
try:
    from PIL import Image
    import pandas as pd
    from docx import Document
    import json
    
    # PDF-specific imports
    from reportlab.lib.pagesizes import A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib import colors
    from PyPDF2 import PdfReader, PdfWriter
    
    # Agno imports
    from agno.tools import Toolkit

except ImportError as e:
    print(f"Missing required package: {e}")
    print("Please install required packages with:")
    print("pip install pillow pandas openpyxl python-docx reportlab PyPDF2 python-magic")
    sys.exit(1)

# Optional imports for video and audio
try:
    from moviepy import VideoFileClip
    VIDEO_SUPPORT = True
except ImportError:
    VIDEO_SUPPORT = False
    print("Warning: moviepy not found. Video conversions will be disabled.")
    print("To enable video features: pip install moviepy and install FFmpeg")

try:
    from pydub import AudioSegment
    AUDIO_SUPPORT = True
except ImportError:
    AUDIO_SUPPORT = False
    print("Warning: pydub not found. Audio conversions will be disabled.")
    print("To enable audio features: pip install pydub")

try:
    import magic
    MAGIC_SUPPORT = True
except ImportError:
    MAGIC_SUPPORT = False
    print("Warning: python-magic not found. File type detection will use extensions only.")
    print("To enable advanced file detection: pip install python-magic")

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

BASE_DIR = r"artifacts/"


class FileConverterToolkit(Toolkit):
    """Universal file converter with advanced PDF features, structured as an Agno Toolkit."""
    
    def __init__(self, **kwargs):
        super().__init__(name="file_converter", tools=[
            self.convert_file,
            self.batch_convert,
            self.merge_pdfs,
            self.split_pdf,
            self.list_supported_formats
        ], **kwargs)
        self.supported_formats = {
            'image': {
                'input': ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp', '.ico'],
                'output': ['.jpg', '.jpeg', '.png', '.bmp', '.gif', '.tiff', '.webp', '.pdf', '.ico']
            },
            'document': {
                'input': ['.txt', '.docx', '.doc', '.rtf', '.html', '.md', '.odt'],
                'output': ['.txt', '.docx', '.pdf', '.rtf', '.html', '.json', '.md']
            },
            'spreadsheet': {
                'input': ['.xlsx', '.xls', '.csv', '.tsv', '.json', '.ods'],
                'output': ['.xlsx', '.xls', '.csv', '.tsv', '.json', '.pdf']
            },
            'pdf': {
                'input': ['.pdf'],
                'output': ['.pdf', '.txt', '.json']
            }
        }
        
        # Add video support if available
        if VIDEO_SUPPORT:
            self.supported_formats['video'] = {
                'input': ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv', '.webm', '.m4v', '.3gp'],
                'output': ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv', '.webm', '.gif', '.m4v', '.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a']
            }
        
        # Add audio support if available
        if AUDIO_SUPPORT:
            self.supported_formats['audio'] = {
                'input': ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a', '.opus'],
                'output': ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a', '.opus']
            }
    
    def detect_file_type(self, file_path: str) -> tuple[Optional[str], str]:
        """Detect file type and extension automatically."""
        try:
            # Get actual file extension
            actual_ext = Path(file_path).suffix.lower()
            
            # Use python-magic to detect MIME type if available
            if MAGIC_SUPPORT:
                mime_type = magic.from_file(file_path, mime=True)
                
                # Map MIME types to extensions
                mime_to_ext = {
                    'image/jpeg': '.jpg',
                    'image/png': '.png',
                    'image/gif': '.gif',
                    'image/bmp': '.bmp',
                    'image/tiff': '.tiff',
                    'image/webp': '.webp',
                    'image/x-icon': '.ico',
                    'video/mp4': '.mp4',
                    'video/avi': '.avi',
                    'video/quicktime': '.mov',
                    'video/x-msvideo': '.avi',
                    'video/webm': '.webm',
                    'audio/mpeg': '.mp3',
                    'audio/wav': '.wav',
                    'audio/flac': '.flac',
                    'audio/aac': '.aac',
                    'audio/ogg': '.ogg',
                    'text/plain': '.txt',
                    'text/html': '.html',
                    'text/markdown': '.md',
                    'application/pdf': '.pdf',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx',
                    'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet': '.xlsx',
                    'application/vnd.ms-excel': '.xls',
                    'text/csv': '.csv',
                    'application/json': '.json',
                }
                
                detected_ext = mime_to_ext.get(mime_type, actual_ext)
                
                # If no extension was provided, use detected extension
                if not actual_ext:
                    actual_ext = detected_ext
            # Find file type category
            file_type = self.get_file_type_by_extension(actual_ext)
            
            return file_type, actual_ext
            
        except Exception as e:
            logger.warning(f"Could not detect file type for {file_path}: {e}")
            # Fallback to extension-based detection
            ext = Path(file_path).suffix.lower()
            return self.get_file_type_by_extension(ext), ext
    
    def get_file_type_by_extension(self, ext: str) -> Optional[str]:
        """Get file type category by extension."""
        for file_type, formats in self.supported_formats.items():
            if ext in formats['input']:
                return file_type
        return None
    
    async def convert_image(self, input_path: str, output_path: str) -> bool:
        """Convert image files."""
        try:
            def _convert():
                with Image.open(input_path) as img:
                    # Convert RGBA to RGB if saving as JPEG
                    if output_path.lower().endswith(('.jpg', '.jpeg')) and img.mode == 'RGBA':
                        img = img.convert('RGB')
                    
                    # Special handling for PDF
                    if output_path.lower().endswith('.pdf'):
                        if img.mode == 'RGBA':
                            img = img.convert('RGB')
                        img.save(output_path, 'PDF')
                    else:
                        img.save(output_path)
                    
                    logger.info(f"Image converted: {input_path} -> {output_path}")
                    return True
            return await asyncio.to_thread(_convert)
        except Exception as e:
            logger.error(f"Error converting image: {e}")
            return False
    
    async def convert_video(self, input_path: str, output_path: str) -> bool:
        """Convert video files."""
        try:
            def _convert():
                video = VideoFileClip(input_path)
                output_ext = Path(output_path).suffix.lower()
                
                # Handle audio extraction from video
                if output_ext in ['.mp3', '.wav', '.flac', '.aac', '.ogg', '.m4a']:
                    audio = video.audio
                    if audio is None:
                        logger.error(f"No audio track found in video: {input_path}")
                        video.close()
                        return False
                    
                    # Extract audio with appropriate format - MoviePy v2 syntax
                    if output_ext == '.mp3':
                        audio.write_audiofile(output_path)
                    elif output_ext == '.wav':
                        audio.write_audiofile(output_path)
                    elif output_ext == '.flac':
                        audio.write_audiofile(output_path, codec='flac')
                    elif output_ext == '.aac':
                        audio.write_audiofile(output_path, codec='aac')
                    elif output_ext == '.ogg':
                        audio.write_audiofile(output_path, codec='libvorbis')
                    elif output_ext == '.m4a':
                        audio.write_audiofile(output_path, codec='aac')
                    
                    audio.close()
                    video.close()
                    logger.info(f"Audio extracted from video: {input_path} -> {output_path}")
                    return True
                
                # Handle video format conversions
                elif output_ext == '.gif':
                    # Resize for smaller file size - MoviePy v2 syntax
                    resized_video = video.resized(height=240)  # Resize to 240p height, maintain aspect ratio
                    resized_video.write_gif(output_path, fps=10)
                    resized_video.close()
                else:
                    video.write_videofile(output_path)
                
                video.close()
                logger.info(f"Video converted: {input_path} -> {output_path}")
                return True
            return await asyncio.to_thread(_convert)
        except Exception as e:
            logger.error(f"Error converting video: {e}")
            return False
    
    async def convert_audio(self, input_path: str, output_path: str) -> bool:
        """Convert audio files."""
        try:
            def _convert():
                audio = AudioSegment.from_file(input_path)
                
                # Get output format
                output_format = Path(output_path).suffix[1:].lower()
                
                # Export with appropriate parameters
                if output_format == 'mp3':
                    audio.export(output_path, format="mp3", bitrate="192k")
                elif output_format == 'wav':
                    audio.export(output_path, format="wav")
                elif output_format == 'flac':
                    audio.export(output_path, format="flac")
                else:
                    audio.export(output_path, format=output_format)
                
                logger.info(f"Audio converted: {input_path} -> {output_path}")
                return True
            return await asyncio.to_thread(_convert)
        except Exception as e:
            logger.error(f"Error converting audio: {e}")
            return False
    
    async def convert_document_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert various document formats to PDF."""
        try:
            def _convert():
                input_ext = Path(input_path).suffix.lower()
                
                # Create PDF document
                doc = SimpleDocTemplate(output_path, pagesize=A4)
                styles = getSampleStyleSheet()
                story = []
                
                # Read content based on input format
                if input_ext == '.txt':
                    with open(input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                        # Split into paragraphs
                        paragraphs = content.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                story.append(Paragraph(para.strip(), styles['Normal']))
                                story.append(Spacer(1, 12))
                
                elif input_ext == '.docx':
                    docx_doc = Document(input_path)
                    for paragraph in docx_doc.paragraphs:
                        if paragraph.text.strip():
                            story.append(Paragraph(paragraph.text, styles['Normal']))
                            story.append(Spacer(1, 12))
                
                elif input_ext == '.html':
                    with open(input_path, 'r', encoding='utf-8') as f:
                        html_content = f.read()
                        # Simple HTML to text conversion (basic)
                        import re
                        text = re.sub('<[^<]+?>', '', html_content)
                        paragraphs = text.split('\n')
                        for para in paragraphs:
                            if para.strip():
                                story.append(Paragraph(para.strip(), styles['Normal']))
                                story.append(Spacer(1, 12))
                
                elif input_ext == '.md':
                    with open(input_path, 'r', encoding='utf-8') as f:
                        md_content = f.read()
                        # Basic markdown to text conversion
                        import re
                        # Remove markdown formatting
                        text = re.sub(r'[#*`_]', '', md_content)
                        paragraphs = text.split('\n\n')
                        for para in paragraphs:
                            if para.strip():
                                story.append(Paragraph(para.strip(), styles['Normal']))
                                story.append(Spacer(1, 12))
                
                # Build PDF
                doc.build(story)
                logger.info(f"Document converted to PDF: {input_path} -> {output_path}")
                return True
            return await asyncio.to_thread(_convert)
        except Exception as e:
            logger.error(f"Error converting document to PDF: {e}")
            return False
    
    async def convert_spreadsheet_to_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert spreadsheet to PDF."""
        try:
            def _convert():
                input_ext = Path(input_path).suffix.lower()
                
                # Read spreadsheet
                if input_ext == '.csv':
                    df = pd.read_csv(input_path)
                elif input_ext == '.tsv':
                    df = pd.read_csv(input_path, sep='\t')
                elif input_ext in ['.xlsx', '.xls']:
                    df = pd.read_excel(input_path)
                elif input_ext == '.json':
                    df = pd.read_json(input_path)
                
                # Create PDF
                doc = SimpleDocTemplate(output_path, pagesize=A4)
                # styles = getSampleStyleSheet()
                story = []
                
                # Convert DataFrame to table data
                table_data = [df.columns.tolist()]  # Header
                for _, row in df.iterrows():
                    table_data.append(row.tolist())
                
                # Create table
                table = Table(table_data)
                table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, 0), 12),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                
                story.append(table)
                doc.build(story)
                
                logger.info(f"Spreadsheet converted to PDF: {input_path} -> {output_path}")
                return True
            return await asyncio.to_thread(_convert)
        except Exception as e:
            logger.error(f"Error converting spreadsheet to PDF: {e}")
            return False
    
    async def convert_document(self, input_path: str, output_path: str) -> bool:
        """Convert document files."""
        try:
            def _convert():
                output_ext = Path(output_path).suffix.lower()
                
                # Handle PDF conversion
                if output_ext == '.pdf':
                    # Note: This needs to be handled differently since we're in a sync context
                    return True  # Will be handled by calling convert_document_to_pdf separately
                
                # Handle other document conversions
                input_ext = Path(input_path).suffix.lower()
                
                # Read input file
                content = ""
                if input_ext == '.txt':
                    with open(input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif input_ext == '.docx':
                    doc = Document(input_path)
                    content = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                elif input_ext == '.html':
                    with open(input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                elif input_ext == '.md':
                    with open(input_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                
                # Write output file
                if output_ext == '.txt':
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(content)
                elif output_ext == '.docx':
                    doc = Document()
                    doc.add_paragraph(content)
                    doc.save(output_path)
                elif output_ext == '.html':
                    html_content = f"""<!DOCTYPE html>
                                        <html>
                                        <head>
                                            <title>Converted Document</title>
                                            <meta charset="UTF-8">
                                        </head>
                                        <body>
                                            <pre>{content}</pre>
                                        </body>
                                        </html>"""
                    with open(output_path, 'w', encoding='utf-8') as f:
                        f.write(html_content)
                elif output_ext == '.json':
                    data = {"content": content, "source": input_path}
                    with open(output_path, 'w', encoding='utf-8') as f:
                        json.dump(data, f, indent=2)
                
                logger.info(f"Document converted: {input_path} -> {output_path}")
                return True
            
            output_ext = Path(output_path).suffix.lower()
            
            # Handle PDF conversion separately since it's async
            if output_ext == '.pdf':
                return await self.convert_document_to_pdf(input_path, output_path)
            
            # Handle other conversions
            return await asyncio.to_thread(_convert)
        except Exception as e:
            logger.error(f"Error converting document: {e}")
            return False
    
    async def convert_spreadsheet(self, input_path: str, output_path: str) -> bool:
        """Convert spreadsheet files."""
        try:
            def _convert():
                output_ext = Path(output_path).suffix.lower()
                
                # Handle PDF conversion
                if output_ext == '.pdf':
                    # Note: This needs to be handled differently since we're in a sync context
                    return True  # Will be handled by calling convert_spreadsheet_to_pdf separately
                
                # Handle other spreadsheet conversions
                input_ext = Path(input_path).suffix.lower()
                
                # Read input file
                if input_ext == '.csv':
                    df = pd.read_csv(input_path)
                elif input_ext == '.tsv':
                    df = pd.read_csv(input_path, sep='\t')
                elif input_ext in ['.xlsx', '.xls']:
                    df = pd.read_excel(input_path)
                elif input_ext == '.json':
                    df = pd.read_json(input_path)
                
                # Write output file
                if output_ext == '.csv':
                    df.to_csv(output_path, index=False)
                elif output_ext == '.tsv':
                    df.to_csv(output_path, sep='\t', index=False)
                elif output_ext == '.xlsx':
                    df.to_excel(output_path, index=False)
                elif output_ext == '.json':
                    df.to_json(output_path, orient='records', indent=2)
                
                logger.info(f"Spreadsheet converted: {input_path} -> {output_path}")
                return True
            return await asyncio.to_thread(_convert)
        except Exception as e:
            logger.error(f"Error converting spreadsheet: {e}")
            return False
    
    async def merge_pdfs(self, input_files: List[str], output_path: str) -> bool:
        """Merge multiple PDF files into one."""
        # NEW: resolve paths
        input_files_resolved = [str(self._resolve_path(p)) for p in input_files]
        output_path_resolved = str(self._resolve_path(output_path))
        try:
            def _merge():
                pdf_writer = PdfWriter()
                for input_file in input_files_resolved:
                    if not os.path.exists(input_file):
                        logger.warning(f"File not found: {input_file}")
                        continue
                    with open(input_file, 'rb') as f:
                        pdf_reader = PdfReader(f)
                        for page in pdf_reader.pages:
                            pdf_writer.add_page(page)
                os.makedirs(os.path.dirname(output_path_resolved), exist_ok=True)
                with open(output_path_resolved, 'wb') as f:
                    pdf_writer.write(f)
                logger.info(f"PDFs merged: {len(input_files_resolved)} files -> {output_path_resolved}")
                return True
            return await asyncio.to_thread(_merge)
        except Exception as e:
            logger.error(f"Error merging PDFs: {e}")
            return False
    
    async def split_pdf(self, input_path: str, output_dir: str, pages_per_file: int = 1) -> bool:
        """Split a PDF into multiple files."""
        # NEW: resolve paths
        input_path_resolved = str(self._resolve_path(input_path))
        output_dir_resolved = str(self._resolve_path(output_dir))
        try:
            def _split():
                if not os.path.exists(input_path_resolved):
                    logger.error(f"Input file not found: {input_path_resolved}")
                    return False
                os.makedirs(output_dir_resolved, exist_ok=True)
                with open(input_path_resolved, 'rb') as f:
                    pdf_reader = PdfReader(f)
                    total_pages = len(pdf_reader.pages)
                    base_name = Path(input_path_resolved).stem
                    for i in range(0, total_pages, pages_per_file):
                        pdf_writer = PdfWriter()
                        for j in range(i, min(i + pages_per_file, total_pages)):
                            pdf_writer.add_page(pdf_reader.pages[j])
                        if pages_per_file == 1:
                            output_filename = f"{base_name}_page_{i + 1}.pdf"
                        else:
                            end_page = min(i + pages_per_file, total_pages)
                            output_filename = f"{base_name}_pages_{i + 1}-{end_page}.pdf"
                        output_file_path = os.path.join(output_dir_resolved, output_filename)
                        with open(output_file_path, 'wb') as output_file:
                            pdf_writer.write(output_file)
                    logger.info(f"PDF split: {input_path_resolved} -> {total_pages} pages into {output_dir_resolved}")
                    return True
            return await asyncio.to_thread(_split)
        except Exception as e:
            logger.error(f"Error splitting PDF: {e}")
            return False
    
    async def extract_pdf_text(self, input_path: str, output_path: str) -> bool:
        """Extract text from PDF."""
        try:
            def _extract():
                with open(input_path, 'rb') as f:
                    pdf_reader = PdfReader(f)
                    text = ""
                    
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    
                    with open(output_path, 'w', encoding='utf-8') as output_file:
                        output_file.write(text)
                    
                    logger.info(f"Text extracted from PDF: {input_path} -> {output_path}")
                    return True
            return await asyncio.to_thread(_extract)
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return False
    
    async def convert_pdf(self, input_path: str, output_path: str) -> bool:
        """Convert PDF to other formats."""
        try:
            def _convert():
                output_ext = Path(output_path).suffix.lower()
                
                if output_ext == '.txt':
                    # Note: This needs to be handled differently since we're in a sync context
                    return True  # Will be handled by calling extract_pdf_text separately
                elif output_ext == '.json':
                    # Extract text and save as JSON
                    with open(input_path, 'rb') as f:
                        pdf_reader = PdfReader(f)
                        pages = []
                        
                        for i, page in enumerate(pdf_reader.pages):
                            pages.append({
                                "page": i + 1,
                                "text": page.extract_text()
                            })
                        
                        data = {
                            "source": input_path,
                            "total_pages": len(pages),
                            "pages": pages
                        }
                        
                        with open(output_path, 'w', encoding='utf-8') as f:
                            json.dump(data, f, indent=2)
                    
                    logger.info(f"PDF converted to JSON: {input_path} -> {output_path}")
                    return True
                
                return False
            
            output_ext = Path(output_path).suffix.lower()
            
            # Handle text extraction separately since it's async
            if output_ext == '.txt':
                return await self.extract_pdf_text(input_path, output_path)
            
            # Handle other conversions
            return await asyncio.to_thread(_convert)
        except Exception as e:
            logger.error(f"Error converting PDF: {e}")
            return False
    
    def _resolve_path(self, p) -> Path:
        """Return absolute Path. If p is relative, prepend BASE_DIR."""
        path = Path(p)
        if not path.is_absolute():
            path = Path(BASE_DIR) / path
        return path

    async def convert_file(self, input_path: str, output_path: str) -> str:
        """
        Convert a single file from one format to another with automatic type detection.

        Args:
            input_path (str): The path to the input file.
            output_path (str): The path to the output file (including new extension).

        Returns:
            str: A message indicating success or failure.
        """
        # NEW: resolve paths
        input_path = str(self._resolve_path(input_path))
        output_path = str(self._resolve_path(output_path))
        if not await asyncio.to_thread(os.path.exists, input_path):
            logger.error(f"Input file not found: {input_path}")
            return f"❌ Error: Input file not found at {input_path}"
        # Create output directory if it doesn't exist
        output_dir = os.path.dirname(output_path)
        if output_dir:
            await asyncio.to_thread(os.makedirs, output_dir, exist_ok=True)
        # Detect file type automatically
        file_type, detected_ext = self.detect_file_type(input_path)
        if not file_type:
            logger.error(f"Unsupported file type: {input_path}")
            return f"❌ Error: Unsupported file type for {input_path}"
        # Validate output format
        output_ext = Path(output_path).suffix.lower()
        if output_ext not in self.supported_formats[file_type]['output']:
            logger.error(f"Unsupported output format {output_ext} for {file_type}")
            return f"❌ Error: Unsupported output format {output_ext} for {file_type} files."
        # Convert based on file type
        success = False
        if file_type == 'image':
            success = await self.convert_image(input_path, output_path)
        elif file_type == 'video':
            success = await self.convert_video(input_path, output_path)
        elif file_type == 'audio':
            success = await self.convert_audio(input_path, output_path)
        elif file_type == 'document':
            success = await self.convert_document(input_path, output_path)
        elif file_type == 'spreadsheet':
            success = await self.convert_spreadsheet(input_path, output_path)
        elif file_type == 'pdf':
            success = await self.convert_pdf(input_path, output_path)
        return f"✅ Successfully converted: {input_path} -> {output_path}" if success else f"❌ Failed to convert: {input_path}"
    
    async def batch_convert(self, input_dir: str, output_dir: str, input_ext: Optional[str] = None, output_ext: Optional[str] = None) -> str:
        """
        Convert multiple files in a directory.

        Args:
            input_dir (str): The directory containing files to convert.
            output_dir (str): The directory where converted files will be saved.
            input_ext (Optional[str]): The specific input extension to filter by (e.g., '.jpg').
            output_ext (Optional[str]): The target output extension (e.g., '.png').

        Returns:
            str: A summary of the batch conversion results.
        """
        # NEW: resolve paths
        input_dir_path = self._resolve_path(input_dir)
        output_dir_path = self._resolve_path(output_dir)
        if not await asyncio.to_thread(input_dir_path.exists):
            logger.error(f"Input directory not found: {input_dir_path}")
            return f"❌ Error: Input directory not found at {input_dir_path}"
        # If no specific extension provided, process all supported files
        def _glob_files():
            if input_ext:
                return list(input_dir_path.glob(f"*{input_ext}"))
            files = []
            for ext_list in [formats['input'] for formats in self.supported_formats.values()]:
                for ext in ext_list:
                    files.extend(input_dir_path.glob(f"*{ext}"))
            return files
        files = await asyncio.to_thread(_glob_files)
        if not files:
            logger.warning(f"No supported files found in {input_dir_path}")
            return f"⚠️ No supported files found in {input_dir_path}"
        successful = 0
        failed = 0
        for file_path in files:
            if output_ext:
                output_file = output_dir_path / f"{file_path.stem}{output_ext}"
            else:
                output_file = output_dir_path / file_path.name
            result = await self.convert_file(str(file_path), str(output_file))
            if "Successfully" in result:
                successful += 1
            else:
                failed += 1
        summary = f"Batch conversion complete: {successful} successful, {failed} failed."
        logger.info(summary)
        return f"✅ {summary}"
    
    def list_supported_formats(self) -> str:
        """
        Display all supported file formats.

        Returns:
            str: A formatted string listing all supported input and output formats.
        """
        result = ["Supported File Formats:"]
        result.append("=" * 50)
        
        for file_type, formats in self.supported_formats.items():
            result.append(f"\n{file_type.upper()}:")
            result.append(f"  Input:  {', '.join(formats['input'])}")
            result.append(f"  Output: {', '.join(formats['output'])}")
        
        return "\n".join(result) 

